"""Tests du rendu Word brut. On relit le .docx et on vérifie le contenu texte,
pas de mise en page (plus de tableaux, plus d'encadré)."""

from docx import Document

from src.output.render_docx import render_docx, render_docx_bytes
from tests.test_render_markdown import make_memo


def test_render_docx_bytes_est_un_docx_valide():
    # Un .docx est une archive ZIP : ses premiers octets sont la signature "PK".
    data = render_docx_bytes(make_memo())
    assert isinstance(data, bytes) and data[:2] == b"PK"
    # Relisible par python-docx, et non vide.
    from io import BytesIO
    doc = Document(BytesIO(data))
    assert len(doc.paragraphs) > 0


def test_render_docx_cree_fichier_et_contenu(tmp_path):
    memo = make_memo()
    path = render_docx(memo, output_dir=tmp_path)

    assert path.exists()
    assert path.name == "memo_Acme_SaaS_2026-07-12.docx"

    doc = Document(str(path))
    assert len(doc.tables) == 0  # rendu brut : aucun tableau Word
    textes = "\n".join(p.text for p in doc.paragraphs)
    assert "Recommandation : APPROFONDIR" in textes
    assert "indisponible" in textes  # bandeau de l'analyse LLM, en texte simple


def test_render_docx_saute_les_separateurs_de_tableau(tmp_path):
    # Les lignes '| --- | --- |' du markdown ne doivent pas apparaître dans le docx.
    path = render_docx(make_memo(), output_dir=tmp_path)
    doc = Document(str(path))
    assert all("---" not in p.text for p in doc.paragraphs)
