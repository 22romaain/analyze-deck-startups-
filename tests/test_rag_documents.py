"""Tests du découpage et de la lecture des documents (RAG, pur, hors ligne)."""

from src.rag.documents import chunk_text, read_documents


def test_chunk_text_rattache_les_sections():
    text = "# Titre A\nphrase un.\n\n## Titre B\nphrase deux.\n"
    chunks = chunk_text(text, "f.md")
    assert [c.section for c in chunks] == ["Titre A", "Titre B"]
    assert chunks[0].text == "phrase un."
    assert all(c.id.startswith("f.md::") for c in chunks)


def test_chunk_text_decoupe_gros_texte():
    # Beaucoup de lignes courtes dans une seule section -> plusieurs passages.
    long = "\n".join(f"ligne {i} avec un peu de texte" for i in range(100))
    chunks = chunk_text(long, "big.txt")
    assert len(chunks) >= 2


def test_read_documents_ignore_readme_et_non_supporte(tmp_path):
    (tmp_path / "README.md").write_text("# instructions", encoding="utf-8")
    (tmp_path / "cours1.md").write_text("# A\ncontenu", encoding="utf-8")
    (tmp_path / "notes.txt").write_text("texte", encoding="utf-8")
    (tmp_path / "image.png").write_bytes(b"x")
    docs = read_documents(tmp_path)
    names = {name for name, _ in docs}
    assert names == {"cours1.md", "notes.txt"}


def test_read_documents_lit_docx_avec_titres(tmp_path):
    # Un .docx est lu, et ses titres Word deviennent des sections comme en Markdown.
    from docx import Document

    doc = Document()
    doc.add_heading("Dimensionner le marche", level=1)
    doc.add_paragraph("Exiger un TAM bottom-up plutot qu un pourcentage.")
    doc.save(tmp_path / "cours.docx")

    texte = dict(read_documents(tmp_path))["cours.docx"]
    assert "# Dimensionner le marche" in texte  # titre -> ligne Markdown
    chunks = chunk_text(texte, "cours.docx")
    assert chunks[0].section == "Dimensionner le marche"
    assert "bottom-up" in chunks[0].text
